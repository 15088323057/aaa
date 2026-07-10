#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""支付系统综合安全漏洞修复报告"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'; s.font.size = Pt(11)
s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.3
s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

PRI = RGBColor(0x1a,0x1a,0x2e); ACC = RGBColor(0x29,0x80,0xB9)
RED = RGBColor(0xE7,0x4C,0x3C); GRN = RGBColor(0x27,0xAE,0x60)
ORG = RGBColor(0xF3,0x9C,0x12); GRY = RGBColor(0x7F,0x8C,0x8D); WHT = RGBColor(0xFF,0xFF,0xFF)

def h(text, level=1, color=PRI):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = color

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    shd = parse_xml('<w:shd {} w:fill="2d2d2d" w:val="clear"/>'.format(nsdecls("w")))
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xF8,0xF8,0xF2)

def box(text, bg="EBF5FB", icon="i", icc=ACC):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    shd = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls("w"), bg))
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(icon+' '); r.font.color.rgb = icc; r.bold = True
    r = p.add_run(text); r.font.size = Pt(10.5)

def shd(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls("w"), color)))

def ct(cell, text, bold=False, color=None, size=None):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text); r.bold = bold
    if color: r.font.color.rgb = color
    if size: r.font.size = size

def tbl(headers, data, widths=None):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        shd(t.rows[0].cells[i],"1a1a2e"); ct(t.rows[0].cells[i],hd,bold=True,color=WHT,size=Pt(10))
    for ri, rd in enumerate(data):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            if ri%2==1: shd(c,"F8F9FA")
            ct(c,str(v),size=Pt(10))
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width = Cm(w)
    doc.add_paragraph('')

# ══ 封面 ══
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
r = p.add_run('='*60); r.font.color.rgb = ACC; r.font.size = Pt(6)
for _ in range(3): doc.add_paragraph('')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('用户信息管理平台'); r.font.size = Pt(32); r.font.color.rgb = PRI; r.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('支付系统安全漏洞综合审计报告'); r.font.size = Pt(20); r.font.color.rgb = ACC
doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('='*40); r.font.color.rgb = ACC; r.font.size = Pt(8)
doc.add_paragraph('')

info = [
    ('审计编号','SEC-PAY-{}-001'.format(datetime.now().strftime("%Y%m%d"))),
    ('审计日期',datetime.now().strftime('%Y年%m月%d日')),
    ('审计范围','/profile 个人中心 + /recharge 充值'),
    ('漏洞总数','6 项（高危4项 / 中危2项）'),
    ('OWASP Top 10','A01:2021 + A03:2021'),
    ('GitHub','github.com/15088323057/aaa'),
]
ti = doc.add_table(rows=len(info),cols=2); ti.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate(info):
    shd(ti.rows[i].cells[0],"1a1a2e"); ct(ti.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(10))
    ti.rows[i].cells[0].width = Cm(4)
    ct(ti.rows[i].cells[1],v,size=Pt(10)); ti.rows[i].cells[1].width = Cm(10)
doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('-- 本报告为安全审计专用文档 --'); r.font.color.rgb = GRY; r.font.size = Pt(9)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(20)
r = p.add_run('='*60); r.font.color.rgb = ACC; r.font.size = Pt(6)
doc.add_page_break()

# ══ 目录 ══
h('目 录',1,PRI); doc.add_paragraph('')
for num,cn in [
    ('1','执行摘要'),('2','审计范围'),
    ('3','漏洞 V-PRO-01：个人中心越权访问 (IDOR)'),
    ('4','漏洞 V-RCH-01：充值接口未认证'),
    ('5','漏洞 V-RCH-02：金额负值注入'),
    ('6','漏洞 V-RCH-03：隐藏字段篡改'),
    ('7','漏洞 V-RCH-04：浮点数精度'),
    ('8','漏洞 V-RCH-05：无充值上限'),
    ('9','修复前后对比'),('10','渗透测试验证'),
    ('11','安全建议'),('12','结论'),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(num+'  '); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PRI
    r = p.add_run(cn); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PRI
doc.add_page_break()

# ══ 1 ══
h('1  执行摘要',1,PRI)
doc.add_paragraph(
    '本报告针对用户管理系统的支付相关功能（个人中心 + 充值）进行综合安全审计。'
    '覆盖了信息泄露、身份认证、输入校验、权限控制、数据精度等维度，'
    '共发现 6 项安全漏洞（高危 4 项，中危 2 项），全部完成修复。')

ct2 = doc.add_table(rows=4,cols=4); ct2.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri,items in enumerate([
    [('漏洞总数','6',ACC),('高危漏洞','4',RED)],
    [('中危漏洞','2',ORG),('修复率','100%',GRN)],
    [('受影响路由','/profile, /recharge',ACC),('CWE 覆盖','5 类',ACC)],
    [('GitHub','已推送',GRN),('安全等级','严重→合规',GRN)],
]):
    for ci,(lb,vl,cl) in enumerate(items):
        c = ct2.rows[ri].cells[ci]; c.text = ''
        p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(vl); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = cl
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(lb); r2.font.size = Pt(9); r2.font.color.rgb = GRY
doc.add_paragraph('')

p = doc.add_paragraph(); r = p.add_run('审计结论：'); r.bold = True
doc.add_paragraph(
    '个人中心存在 IDOR 漏洞，任何用户可通过修改 URL 参数查看他人敏感信息。'
    '充值功能存在 5 项漏洞。所有漏洞已通过 session 统一认证、'
    '服务端权限校验、金额边界检查等方式完成修复。')
doc.add_page_break()

# ══ 2 ══
h('2  审计范围',1,PRI)
doc.add_heading('2.1 功能范围',2)
tbl(['功能','路由','漏洞数','严重程度'],[
    ['个人中心查看','GET /profile','1','高危'],
    ['用户充值','POST /recharge','5','高危/中危'],
],[3,3.5,2.5,3.5])

doc.add_heading('2.2 攻击面模型',2)
doc.add_paragraph('攻击者可通过以下链路实施攻击：',style='List Bullet')
doc.add_paragraph('步骤1：利用 profile IDOR 遍历所有用户的 user_id，收集用户信息',style='List Bullet')
doc.add_paragraph('步骤2：利用 recharge 未认证 + 负值注入，清空目标用户余额',style='List Bullet')
doc.add_paragraph('步骤3：利用 recharge 无上限 + 隐藏字段篡改，给自己充值任意金额',style='List Bullet')
doc.add_page_break()

# ══ 3 ══
h('3  漏洞 V-PRO-01：个人中心越权访问 (IDOR)',1,RED)
doc.add_heading('3.1 漏洞描述',2)
doc.add_paragraph(
    '个人中心从未认证用户处获取 user_id 参数，直接查询并展示用户资料。'
    '攻击者可通过遍历 user_id 参数（1,2,3...）获取所有用户的邮箱、手机号、余额等敏感信息。'
    '这是典型的 IDOR（不安全的直接对象引用）漏洞。'
)
it = doc.add_table(rows=5,cols=2); it.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('漏洞编号','V-PRO-01'),('CWE','CWE-639: Authorization Bypass Through User-Controlled Key'),
    ('CVSS 3.1','7.5 (HIGH)'),('OWASP','A01:2021 权限控制失效'),
    ('泄露信息','邮箱、手机号、余额')]):
    shd(it.rows[i].cells[0],"1a1a2e"); ct(it.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(9))
    ct(it.rows[i].cells[1],v,size=Pt(9))
doc.add_paragraph('')

doc.add_heading('3.2 问题代码',2)
code(
    '@app.route("/profile")\n'
    'def profile():\n'
    '    user_id = request.args.get("user_id")  # 从URL参数获取\n'
    '    user = get_user_by_id(user_id)          # 直接查询\n'
    '    return render_template("profile.html", user=user)\n'
    '    # 没有校验这个user_id是不是当前用户的！'
)

doc.add_heading('3.3 攻击场景',2)
code(
    '# 攻击者登录后，遍历所有user_id\n'
    'curl http://target/profile?user_id=1  # 查看 admin: 99999元\n'
    'curl http://target/profile?user_id=2  # 查看 alice: 100元\n'
    'curl http://target/profile?user_id=3  # 查看 bob: 500元\n'
    '# 批量收集所有用户敏感信息'
)

doc.add_heading('3.4 修复方案',2)
box('不再从 URL 参数获取 user_id，改为从 session 获取当前登录用户。',"E8F5E9","✅",GRN)
code(
    '@app.route("/profile")\n'
    'def profile():\n'
    '    username = session.get("username")  # 从session获取\n'
    '    user = get_user_by_username(username)  # 只能查自己\n'
    '    return render_template("profile.html", user=user)\n'
    '    # 即使URL加了 ?user_id=2 也无效，被完全忽略'
)
box('用户传入的 user_id 参数被完全忽略，从源头上杜绝了IDOR漏洞。',"EBF5FB","ℹ",ACC)
doc.add_page_break()

# ══ 4-8 (shorter versions combining the recharge fixes into one V-RCH category) ══
h('4  漏洞 V-RCH-01：充值接口未认证',1,RED)
doc.add_paragraph('充值接口未要求登录，任何人可直接调用。')
code(
    '# 修复前\n'
    'def recharge():\n'
    '    ...  # 没有session检查\n\n'
    '# 修复后\n'
    'def recharge():\n'
    '    session_username = session.get("username")\n'
    '    if not session_username:\n'
    '        return redirect("/login")'
)

h('5  漏洞 V-RCH-02：金额负值注入',1,RED)
doc.add_paragraph('amount 未校验正负，负值相当于扣款。')
code(
    '# 修复前\n'
    'amount = float(amount)  # amount=-99999 直接通过\n\n'
    '# 修复后\n'
    'amount = round(float(amount), 2)\n'
    'if amount <= 0:\n'
    '    return render_template("profile.html", error="充值金额必须大于零")'
)

h('6  漏洞 V-RCH-03：隐藏字段篡改',1,RED)
doc.add_paragraph(
    '充值表单使用隐藏字段传递 user_id，攻击者可修改该值给他人充值。'
    '修复后充值目标由 session 自动确定，不再依赖前端传入的 user_id。')
code(
    '# 修复前: 依赖前端隐藏字段\n'
    '<input type="hidden" name="user_id" value="{{ user.id }}">\n\n'
    '# 修复后: 完全由服务端session决定\n'
    'current_user = get_user_by_username(session_username)  # 只能给自己充'
)

h('7  漏洞 V-RCH-04：浮点数精度问题',1,ORG)
doc.add_paragraph(
    '直接使用 float 存储金额，微小误差在多次交易后会累积。'
    '修复后使用 round() 保留 2 位小数。')
code('# 修复前\namount = float(amount)\n\n# 修复后\namount = round(float(amount), 2)')

h('8  漏洞 V-RCH-05：无充值上限',1,ORG)
doc.add_paragraph('无单次充值上限，可充任意金额。修复后上限为 999,999 元。')
code('# 修复后\nif amount > 999999:\n    return render_template("profile.html", error="单次充值金额不能超过 999,999 元")')
doc.add_page_break()

# ══ 9 ══
h('9  修复前后对比',1,PRI)
doc.add_heading('9.1 /profile 路由对比',2)
tbl(['维度','修复前（有漏洞）','修复后（安全）'],[
    ['用户身份来源','URL 参数 ?user_id=X','session.get("username")'],
    ['越权访问','可查看任意用户资料','只能查看自己的资料'],
    ['遍历攻击','user_id=1,2,3... 可遍历','参数被忽略，无法攻击'],
    ['未登录访问','返回用户资料','重定向到 /login'],
    ['攻击难度','极低（改 URL 即可）','攻击无效'],
],[3,5.5,5.5])

doc.add_heading('9.2 /recharge 路由对比',2)
tbl(['维度','修复前（有漏洞）','修复后（安全）'],[
    ['登录要求','无要求','必须登录'],
    ['金额正负','可正可负','必须大于零'],
    ['充值对象','隐藏字段指定','session 自动确定'],
    ['金额上限','无上限','999,999 元'],
    ['资金精度','float 直接存储','round() 保留2位'],
    ['攻击路径','5 条','0 条'],
],[3,5.5,5.5])
doc.add_page_break()

# ══ 10 ══
h('10  渗透测试验证',1,PRI)
tbl(['编号','测试场景','修复前结果','修复后结果','状态'],[
    ['T01','profile?user_id=2 越权访问','可看到alice资料','只能看到自己的','✅'],
    ['T02','profile?user_id=999 遍历','返回用户资料','重定向到 /login','✅'],
    ['T03','未登录调用 recharge','充值成功','重定向到 /login','✅'],
    ['T04','amount=-99999 负值注入','扣款成功','拒绝：金额须大于零','✅'],
    ['T05','修改隐藏字段user_id','给他人充值','拒绝：只能充自己','✅'],
    ['T06','amount=0 零值','充值成功（余额不变）','拒绝：金额须大于零','✅'],
    ['T07','amount=9999999 超上限','充值成功','拒绝：超上限','✅'],
    ['T08','amount=0.01 最小正数','成功','成功（正常业务）','✅'],
],[1.5,4,3.5,3.5,1.5])
doc.add_page_break()

# ══ 11 ══
h('11  安全建议',1,PRI)
for t,d in [
    ('11.1  统一权限装饰器','实现 @login_required 装饰器，统一应用到所有需要认证的路由。'),
    ('11.2  金融数据用整数分','将金额单位改为"分"，数据库存整数（如 10000 分=100元），彻底避免浮点数问题。'),
    ('11.3  交易流水表','新增 transactions 表记录每笔充值：时间、用户、金额、IP。便于审计和异常检测。'),
    ('11.4  频率限制','同一用户每分钟最多充值 3 次，阻止自动化脚本。'),
    ('11.5  HTTPS 部署','生产环境必须配置 HTTPS，防止中间人攻击篡改请求。'),
    ('11.6  定期代码审计','使用 Bandit、SonarQube 自动扫描，及时发现漏洞。'),
]:
    doc.add_heading(t,2); doc.add_paragraph(d)
doc.add_page_break()

# ══ 12 ══
h('12  结论',1,PRI)
doc.add_paragraph(
    '本次审计覆盖个人中心和充值两个支付相关功能，共发现 6 项安全漏洞。'
    '其中最严重的是个人中心 IDOR 漏洞和充值金额负值注入漏洞，'
    '组合利用可导致任意用户余额被篡改。')
doc.add_paragraph(
    '所有漏洞已修复。核心改进包括：所有敏感操作绑定 session 认证、'
    '服务端决定操作权限而非依赖前端参数、金额正负及上限校验。'
    '系统支付安全等级已从严重风险提升至合规水平。')

doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('='*40); r.font.color.rgb = ACC; r.font.size = Pt(8)
doc.add_paragraph('')

st = doc.add_table(rows=6,cols=2); st.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('审计功能','/profile + /recharge'),
    ('发现漏洞','6 项（高危 4 / 中危 2）'),
    ('已修复','6 项（修复率 100%）'),
    ('修复方式','session 统一认证 / 服务端权限 / 金额校验'),
    ('安全等级','严重风险 -> 安全合规'),
    ('报告版本','v1.0 | {}'.format(datetime.now().strftime('%Y-%m-%d'))),
]):
    shd(st.rows[i].cells[0],"1a1a2e"); ct(st.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(10))
    ct(st.rows[i].cells[1],v,size=Pt(10))

doc.add_paragraph(''); doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('-- 报告完 --'); r.font.size = Pt(12); r.font.color.rgb = GRY; r.italic = True

out = '/workspace/支付系统安全漏洞综合审计报告.docx'
doc.save(out)
print('报告生成: ' + out)
print('大小: {:.1f} KB'.format(os.path.getsize(out)/1024))
