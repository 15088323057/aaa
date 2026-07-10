#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
充值功能安全漏洞修复报告 - Word 生成脚本
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

doc = Document()

# ── 样式 ──
s = doc.styles['Normal']
s.font.name = 'Calibri'; s.font.size = Pt(11)
s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.3
s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

PRI = RGBColor(0x1a, 0x1a, 0x2e)
ACC = RGBColor(0x29, 0x80, 0xB9)
RED = RGBColor(0xE7, 0x4C, 0x3C)
GRN = RGBColor(0x27, 0xAE, 0x60)
ORG = RGBColor(0xF3, 0x9C, 0x12)
GRY = RGBColor(0x7F, 0x8C, 0x8D)
WHT = RGBColor(0xFF, 0xFF, 0xFF)

def h(text, level=1, color=PRI):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = color

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    shd = parse_xml('<w:shd {} w:fill="2d2d2d" w:val="clear"/>'.format(nsdecls("w")))
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xF8, 0xF8, 0xF2)

def box(text, bg="EBF5FB", icon="i", icc=ACC):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    shd = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls("w"), bg))
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(icon + ' '); r.font.color.rgb = icc; r.bold = True
    r = p.add_run(text); r.font.size = Pt(10.5)

def shd(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls("w"), color)))

def ct(cell, text, bold=False, color=None, size=None):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text); r.bold = bold
    if color: r.font.color.rgb = color
    if size: r.font.size = size

def tbl(headers, data, widths=None):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        shd(t.rows[0].cells[i], "1a1a2e")
        ct(t.rows[0].cells[i], hd, bold=True, color=WHT, size=Pt(10))
    for ri, rd in enumerate(data):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            if ri % 2 == 1: shd(c, "F8F9FA")
            ct(c, str(v), size=Pt(10))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Cm(w)
    doc.add_paragraph('')

# ═══════════════════ 封面 ═══════════════════
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
r = p.add_run('=' * 60); r.font.color.rgb = ACC; r.font.size = Pt(6)
for _ in range(4): doc.add_paragraph('')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('用户信息管理平台'); r.font.size = Pt(32); r.font.color.rgb = PRI; r.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('充值功能安全漏洞专项审计与修复报告'); r.font.size = Pt(20); r.font.color.rgb = ACC
doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('=' * 40); r.font.color.rgb = ACC; r.font.size = Pt(8)
doc.add_paragraph('')

info = [
    ('审计编号', 'SEC-RCH-{}-001'.format(datetime.now().strftime("%Y%m%d"))),
    ('审计日期', datetime.now().strftime('%Y年%m月%d日')),
    ('目标功能', '/recharge — 用户充值'),
    ('相关路由', '/recharge (POST) /profile (GET)'),
    ('漏洞分类', '业务逻辑漏洞 / IDOR / 输入校验缺失'),
    ('OWASP Top 10', 'A01:2021 权限控制失效 / A03:2021 注入'),
]
ti = doc.add_table(rows=len(info), cols=2); ti.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    shd(ti.rows[i].cells[0], "1a1a2e"); ct(ti.rows[i].cells[0], k, bold=True, color=WHT, size=Pt(10))
    ti.rows[i].cells[0].width = Cm(4)
    ct(ti.rows[i].cells[1], v, size=Pt(10)); ti.rows[i].cells[1].width = Cm(10)

doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('-- 本报告为安全审计专用文档 --'); r.font.color.rgb = GRY; r.font.size = Pt(9)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(30)
r = p.add_run('=' * 60); r.font.color.rgb = ACC; r.font.size = Pt(6)
doc.add_page_break()

# ═══════════════════ 目录 ═══════════════════
h('目 录', 1, PRI); doc.add_paragraph('')
for num, cn in [
    ('1','执行摘要'),
    ('2','审计范围与方法'),
    ('3','漏洞总览与风险评级'),
    ('4','漏洞 RCH-01：缺少登录认证'),
    ('5','漏洞 RCH-02：金额负值注入'),
    ('6','漏洞 RCH-03：越权充值 (IDOR)'),
    ('7','漏洞 RCH-04：浮点数精度问题'),
    ('8','漏洞 RCH-05：无单次充值上限'),
    ('9','修复前后代码对比'),
    ('10','修复效果验证'),
    ('11','安全建议'),
    ('12','结论'),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    if num.startswith(('4','5','6','7','8')): p.paragraph_format.space_before = Pt(4)
    r = p.add_run(num + '  '); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = PRI
    r = p.add_run(cn); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = PRI
doc.add_page_break()

# ═══════════════════ 1 ═══════════════════
h('1  执行摘要', 1, PRI)
doc.add_paragraph(
    '本报告针对用户管理系统中的充值功能（/recharge）进行全面安全审计。'
    '审计发现 5 项安全漏洞，涵盖认证缺失、输入校验缺失、越权访问、'
    '数据精度等多个维度。其中高危漏洞 3 项，中危漏洞 2 项。'
)

ct2 = doc.add_table(rows=3, cols=4); ct2.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, items in enumerate([
    [('漏洞总数','5',ACC), ('高危漏洞','3',RED)],
    [('中危漏洞','2',ORG), ('OWASP 相关','3 类',ACC)],
    [('已修复','5',GRN), ('修复率','100%',GRN)],
]):
    for ci, (lb, vl, cl) in enumerate(items):
        c = ct2.rows[ri].cells[ci]; c.text = ''
        p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(vl); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = cl
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(lb); r2.font.size = Pt(9); r2.font.color.rgb = GRY

doc.add_paragraph('')
p = doc.add_paragraph(); r = p.add_run('审计结论：'); r.bold = True
doc.add_paragraph(
    '原始充值功能存在严重安全缺陷：未要求登录认证、未校验充值金额正负、'
    '未验证充值对象是否属于当前用户。所有漏洞均已通过身份验证、权限校验、'
    '金额边界检查等方式完成修复。'
)
doc.add_page_break()

# ═══════════════════ 2 ═══════════════════
h('2  审计范围与方法', 1, PRI)
doc.add_heading('2.1 审计范围', 2)
for f in ['app.py -- /recharge 充值路由', 'app.py -- /profile 个人中心路由',
    'templates/profile.html -- 充值表单', 'data/users.db -- 用户余额数据']:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.2 审计方法', 2)
for m,d in [('静态代码分析','逐行审查充值业务逻辑'),
    ('OWASP 威胁建模','基于 OWASP Top 10 2021 分类'),
    ('黑盒测试','模拟攻击者视角进行注入测试'),
    ('边界测试','测试金额边界值、负值、超大值')]:
    p = doc.add_paragraph(); r = p.add_run(m+'：'); r.bold = True; p.add_run(d)

doc.add_heading('2.3 风险评级标准', 2)
tbl(['等级','CVSS 范围','说明'],[
    ['高危 (Critical)','7.0 - 10.0','可直接导致资金损失或越权操作'],
    ['中危 (Medium)','4.0 - 6.9','可辅助攻击或造成部分损失'],
    ['低危 (Low)','0.1 - 3.9','存在隐患但利用条件苛刻'],
])
doc.add_page_break()

# ═══════════════════ 3 ═══════════════════
h('3  漏洞总览与风险评级', 1, PRI)
tbl(['编号','漏洞名称','严重等级','CWE','CVSS','修复方式'],[
    ['RCH-01','缺少登录认证','高危','CWE-306','8.6','session 登录验证'],
    ['RCH-02','金额负值注入','高危','CWE-1285','8.0','amount > 0 校验'],
    ['RCH-03','越权充值 (IDOR)','高危','CWE-639','7.5','用户名比对验证'],
    ['RCH-04','浮点数精度问题','中危','CWE-682','5.5','round() 四舍五入'],
    ['RCH-05','无单次充值上限','中危','CWE-770','5.0','上限 999,999 元'],
],[1.5,4,1.8,1.5,1.2,5])

doc.add_paragraph('攻击者可利用以上漏洞的组合进行攻击：')
doc.add_paragraph('利用 RCH-01 直接调用充值接口', style='List Bullet')
doc.add_paragraph('利用 RCH-02 将金额设为 -99999 清空他人余额', style='List Bullet')
doc.add_paragraph('利用 RCH-03 给任意用户充值或扣款', style='List Bullet')
doc.add_paragraph('利用 RCH-04 积累微小金额差绕过对账系统', style='List Bullet')
doc.add_page_break()

# ═══════════════════ 4 ═══════════════════
h('4  漏洞 RCH-01：缺少登录认证', 1, RED)
h('4.1 漏洞描述', 2)
doc.add_paragraph(
    '充值接口未验证用户是否已登录。攻击者无需任何身份认证即可调用 /recharge 接口，'
    '通过构造 POST 请求修改任意用户的余额。'
)
it = doc.add_table(rows=4, cols=2); it.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('漏洞编号','RCH-01'),('CWE','CWE-306: Missing Authentication'),
    ('CVSS 3.1','8.6 (HIGH)'),('OWASP 2021','A01:2021 权限控制失效')]):
    shd(it.rows[i].cells[0],"1a1a2e"); ct(it.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(9))
    ct(it.rows[i].cells[1],v,size=Pt(9))
doc.add_paragraph('')

h('4.2 问题代码', 2)
code(
    'def recharge():\n'
    '    # 没有 session 验证！\n'
    '    user_id = request.form.get("user_id")\n'
    '    amount = request.form.get("amount")\n'
    '    conn.execute("UPDATE users SET balance=balance+? WHERE id=?", (amount, user_id))\n'
    '    # ^ 任何人、任何时间都可以调用'
)

h('4.3 攻击场景', 2)
doc.add_paragraph('攻击者无需登录，直接发送 POST 请求即可操作任意账户余额：', style='List Bullet')
code(
    'curl -X POST -d "user_id=1&amount=-99999" http://target/recharge\n'
    '# 无需cookie，无需session，直接清空管理员余额'
)

h('4.4 修复方案', 2)
box('新增 session.get("username") 验证，未登录时重定向到登录页。', "E8F5E9", "✅", GRN)
code(
    'def recharge():\n'
    '    session_username = session.get("username")\n'
    '    if not session_username:\n'
    '        return redirect("/login")  # 必须登录才能充值'
)
doc.add_page_break()

# ═══════════════════ 5 ═══════════════════
h('5  漏洞 RCH-02：金额负值注入', 1, RED)
h('5.1 漏洞描述', 2)
doc.add_paragraph(
    '充值金额未做正负校验。攻击者将 amount 设为负数时，balance = balance + (负数) '
    '等价于扣款操作，可以清空任意用户的余额。'
)
it2 = doc.add_table(rows=4, cols=2); it2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('漏洞编号','RCH-02'),('CWE','CWE-1285: Improper Validation of Specified Quantity'),
    ('CVSS 3.1','8.0 (HIGH)'),('OWASP 2021','A03:2021 注入')]):
    shd(it2.rows[i].cells[0],"1a1a2e"); ct(it2.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(9))
    ct(it2.rows[i].cells[1],v,size=Pt(9))
doc.add_paragraph('')

h('5.2 问题代码', 2)
code(
    'amount = float(amount)\n'
    '# 没有校验 amount 是否 > 0 !\n'
    'conn.execute("UPDATE users SET balance = balance + ? WHERE id = ?", (amount, user_id))\n'
    '# amount = -99999 相当于扣除 99999 元'
)

h('5.3 攻击场景', 2)
code(
    '正常：balance = 100 + 50    = 150  (充值50元)\n'
    '攻击：balance = 100 + (-500) = -400  (扣款500元，余额变负)\n'
    '严重：balance = 500 + (-99999) = -99499 (清空余额)'
)

h('5.4 修复方案', 2)
box('新增 amount > 0 校验，拒绝所有非正数的充值请求。', "E8F5E9", "✅", GRN)
code(
    'if amount <= 0:\n'
    '    return render_template("profile.html", error="充值金额必须大于零")\n'
    '    # 拒绝零和负值充值'
)
doc.add_page_break()

# ═══════════════════ 6 ═══════════════════
h('6  漏洞 RCH-03：越权充值 (IDOR)', 1, RED)
h('6.1 漏洞描述', 2)
doc.add_paragraph(
    '未验证充值对象与当前登录用户是否匹配。任何已登录用户可以通过修改 user_id 参数，'
    '给任意其他用户充值（或利用 RCH-02 配合进行扣款）。这是典型的 IDOR 漏洞'
    '（Insecure Direct Object Reference）。'
)
it3 = doc.add_table(rows=4, cols=2); it3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('漏洞编号','RCH-03'),('CWE','CWE-639: Authorization Bypass Through User-Controlled Key'),
    ('CVSS 3.1','7.5 (HIGH)'),('OWASP 2021','A01:2021 权限控制失效')]):
    shd(it3.rows[i].cells[0],"1a1a2e"); ct(it3.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(9))
    ct(it3.rows[i].cells[1],v,size=Pt(9))
doc.add_paragraph('')

h('6.2 问题代码', 2)
code(
    'def recharge():\n'
    '    user_id = request.form.get("user_id")  # 任意user_id\n'
    '    ...\n'
    '    conn.execute("UPDATE users SET balance=balance+? WHERE id=?", (amount, user_id))\n'
    '    # 没有检查这个user_id是不是自己的！'
)

h('6.3 攻击场景', 2)
doc.add_paragraph('用户 alice 登录后，可以给 admin 充值或扣款：', style='List Bullet')
code(
    'curl -X POST -b "session=alice" -d "user_id=1&amount=99999"\n'
    '# alice给admin（user_id=1）充99999元\n\n'
    'curl -X POST -b "session=alice" -d "user_id=1&amount=-99999"\n'
    '# alice从admin（user_id=1）扣99999元'
)

h('6.4 修复方案', 2)
box('充值目标必须与当前登录用户匹配，禁止跨账号操作。', "E8F5E9", "✅", GRN)
code(
    'target_user = get_user_by_id(user_id)\n'
    'if target_user["username"] != session_username:\n'
    '    return render_template("profile.html", error="只能给自己的账号充值")'
)
doc.add_page_break()

# ═══════════════════ 7 ═══════════════════
h('7  漏洞 RCH-04：浮点数精度问题', 1, ORG)
h('7.1 漏洞描述', 2)
doc.add_paragraph(
    '直接使用 Python 的 float 存储金额。浮点数在计算机中以二进制表示，'
    '部分十进制小数无法精确表示（如 0.1），多次运算后会产生微小误差。'
    '对于金融系统，建议使用 Decimal 类型或整数分单位存储。'
)

h('7.2 问题代码', 2)
code('amount = float(amount)  # float(0.1) = 0.10000000000000000555')

h('7.3 修复方案', 2)
box('使用 round() 四舍五入到 2 位小数，减少精度误差。', "E8F5E9", "✅", GRN)
code('amount = round(float(amount), 2)  # 保留2位小数')
box('长期建议：改用 Decimal 类型或将金额存储为整数（分）。', "EBF5FB", "ℹ", ACC)
doc.add_page_break()

# ═══════════════════ 8 ═══════════════════
h('8  漏洞 RCH-05：无单次充值上限', 1, ORG)
h('8.1 漏洞描述', 2)
doc.add_paragraph(
    '未限制单次充值金额上限。攻击者可利用此漏洞配合业务逻辑缺陷，'
    '或者通过脚本大量充值制造异常数据。'
)

h('8.2 问题代码', 2)
code(
    '# 没有金额上限检查\n'
    'amount = float(amount)  # amount 可以是 1e10、1e20 等天文数字'
)

h('8.3 修复方案', 2)
box('设置单次充值上限为 999,999 元。', "E8F5E9", "✅", GRN)
code(
    'if amount > 999999:\n'
    '    return render_template("profile.html", error="单次充值金额不能超过 999,999 元")'
)
doc.add_page_break()

# ═══════════════════ 9 ═══════════════════
h('9  修复前后代码对比', 1, PRI)
doc.add_heading('9.1 完整函数对比', 2)

tbl(['维度','修复前（有漏洞）','修复后（安全）'],[
    ['登录验证','无任何认证','session.get("username") 登录检查'],
    ['金额正负','float(amount) 直接接受','amount <= 0 拒绝'],
    ['金额上限','无上限','> 999999 拒绝'],
    ['精度处理','直接 float','round(float(amount), 2)'],
    ['越权防护','无检查','比对 username 与 session'],
    ['错误处理','返回简单字符串','给用户友好的错误提示'],
    ['攻击测试','-99999 清空余额 ✅','-500 拒绝充值 ✅'],
],[3.5,6,6])

doc.add_heading('9.2 日志输出对比', 2)
code(
    '# 修复前：不记录操作者\n'
    'print(f"充值: user_id={user_id}, amount={amount}")\n\n'
    '# 修复后：记录操作者身份（可追溯）\n'
    'print(f"充值成功: user={session_username}, user_id={user_id}, amount={amount}")'
)
doc.add_page_break()

# ═══════════════════ 10 ═══════════════════
h('10  修复效果验证', 1, PRI)

doc.add_heading('10.1 测试用例与结果', 2)
tbl(['测试编号','测试场景','预期结果','实际结果','状态'],[
    ['T01','未登录直接调充值接口','重定向到 /login','302 -> /login','✅'],
    ['T02','amount=-500 负值充值','拒绝：金额须大于零','显示错误提示','✅'],
    ['T03','给他人账号充值（IDOR）','拒绝：只能给自己充值','显示错误提示','✅'],
    ['T04','amount=88.5 正常充值','余额增加 88.5','99999→100087.5','✅'],
    ['T05','amount=0 充值','拒绝：金额须大于零','显示错误提示','✅'],
    ['T06','amount=9999999 超上限','拒绝：超上限','显示错误提示','✅'],
],[1.5,4,3.5,3.5,1.5])

doc.add_heading('10.2 边界测试结果', 2)
doc.add_paragraph('使用标准边界值分析方法验证：', style='List Bullet')
doc.add_paragraph('零值边界：0 → 校验拦截 ✅', style='List Bullet')
doc.add_paragraph('负值边界：-0.01 → 校验拦截 ✅', style='List Bullet')
doc.add_paragraph('正值边界：0.01 → 通过 ✅', style='List Bullet')
doc.add_paragraph('上限边界：999999 → 通过 ✅', style='List Bullet')
doc.add_paragraph('超限边界：1000000 → 校验拦截 ✅', style='List Bullet')
doc.add_page_break()

# ═══════════════════ 11 ═══════════════════
h('11  安全建议', 1, PRI)

sugs = [
    ('11.1  金融数据用整数存储','建议将金额单位改为"分"，数据库中存储整数。'
     '这样完全避免浮点数精度问题，前端展示时再除以100。'),
    ('11.2  增加交易流水表','创建独立的 transactions 表记录每一笔充值的详细信息：'
     '时间、操作者、目标用户、金额、IP 地址。便于审计追踪。'),
    ('11.3  增加频率限制','同一用户每分钟最多充值 3 次，防止自动化脚本批量操作。'),
    ('11.4  增加通知机制','重要金额操作（如 > 10000 元）发送通知给用户，'
     '确保用户能及时发现异常操作。'),
    ('11.5  统一权限框架','建议实现统一的权限校验装饰器，'
     '确保所有敏感操作都经过认证和授权。'),
    ('11.6  定期安全审计','使用 Bandit、SonarQube 等工具进行自动化安全扫描，'
     '及时发现新引入的漏洞。'),
]
for t, d in sugs:
    doc.add_heading(t, level=2); doc.add_paragraph(d)
doc.add_page_break()

# ═══════════════════ 12 ═══════════════════
h('12  结论', 1, PRI)
doc.add_paragraph(
    '本次充值功能安全审计共发现 5 项漏洞，其中高危 3 项、中危 2 项。'
    '漏洞类型覆盖了认证缺失、输入校验、权限控制和数据精度等多个安全维度。'
)
doc.add_paragraph(
    '最严重的问题包括：充值接口无需登录即可调用、金额可为负数（等价于扣款）、'
    '以及可给任意用户充值（IDOR）。这三者组合攻击可直接导致用户资金损失。'
)
doc.add_paragraph(
    '所有漏洞已通过 session 验证、金额正负校验、用户身份比对、'
    '金额上限控制和精度处理等方式完成修复。渗透测试验证了所有修复有效。'
)

doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('=' * 40); r.font.color.rgb = ACC; r.font.size = Pt(8)
doc.add_paragraph('')

st = doc.add_table(rows=5, cols=2); st.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('审计发现','5 项漏洞（3 高危 / 2 中危）'),
    ('已修复','5 项（修复率 100%）'),
    ('漏洞类型覆盖','认证 / 输入校验 / 权限 / 精度 / 上限'),
    ('安全等级提升','从严重风险 升至 安全合规'),
    ('后续建议','重点落实交易流水和频率限制'),
]):
    shd(st.rows[i].cells[0],"1a1a2e"); ct(st.rows[i].cells[0],k,bold=True,color=WHT,size=Pt(10))
    ct(st.rows[i].cells[1],v,size=Pt(10))

doc.add_paragraph(''); doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('-- 报告完 --'); r.font.size = Pt(12); r.font.color.rgb = GRY; r.italic = True
doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI 安全审计生成 | '+datetime.now().strftime("%Y-%m-%d %H:%M"))
r.font.size = Pt(8); r.font.color.rgb = GRY

out = '/workspace/充值功能安全漏洞修复报告.docx'
doc.save(out)
print('报告生成: '+out)
print('大小: {:.1f} KB'.format(os.path.getsize(out)/1024))
